"""
agents/report_agent.py — Export reports as PDF, XLSX, DOCX.

Reads data from Google Sheets and builds formatted documents.
Uploads each file to Google Shared Drive and returns the share links.
"""
import io
from datetime import datetime
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
import openpyxl
from docx import Document as DocxDoc
import google.auth
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
from backend.services.sheets import SheetsService
from backend.config import GOOGLE_CREDS_JSON, SHARED_DRIVE_ID

_DRIVE_SCOPES = ["https://www.googleapis.com/auth/drive"]


def _drive_credentials():
    if GOOGLE_CREDS_JSON:
        import json
        from google.oauth2.service_account import Credentials
        return Credentials.from_service_account_info(
            json.loads(GOOGLE_CREDS_JSON), scopes=_DRIVE_SCOPES)
    creds, _ = google.auth.default(scopes=_DRIVE_SCOPES)
    if hasattr(creds, "refresh"):
        creds.refresh(Request())
    return creds

# Map report_type → (sheet_tab, column_headers)
REPORT_CONFIG = {
    "attendance":     ("attendance",   ["Date","Lecture No","Subject","PRN","Name","Status","Scan Time","GPS","QR"]),
    "quiz":           ("quiz",         ["Date","Lecture","Subject","Topic","PRN","Name","Score","Total","%","Time(s)"]),
    "teaching_plan":  ("teaching_plan",["Unit","Topic No","Topic","Planned Date","Actual Date","Planned Hrs","Done Hrs","Remaining","CO","Bloom's","Status"]),
    "question_bank":  ("question_bank",["ID","Subject","Unit","Topic","Difficulty","Bloom's","CO","Type","Question","Answer","Explanation"]),
}


def _upload_to_drive(filename: str, mime: str, buf: io.BytesIO) -> str:
    svc   = build("drive", "v3", credentials=_drive_credentials())
    meta  = {"name": filename, "parents": [SHARED_DRIVE_ID]}
    media = MediaIoBaseUpload(buf, mimetype=mime, resumable=False)
    file  = svc.files().create(
        body=meta, media_body=media,
        fields="id,webViewLink", supportsAllDrives=True,
    ).execute()
    return file.get("webViewLink", "")


def _build_pdf(title: str, headers: list, rows: list) -> io.BytesIO:
    buf  = io.BytesIO()
    doc  = SimpleDocTemplate(buf, pagesize=A4)
    s    = getSampleStyleSheet()
    elems = [
        Paragraph(f"<b>{title}</b>", s["Title"]),
        Paragraph(datetime.now().strftime("Generated: %Y-%m-%d %H:%M"), s["Normal"]),
        Spacer(1, 10),
    ]
    data = [headers] + rows
    tbl  = Table(data, repeatRows=1)
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#1e40af")),
        ("TEXTCOLOR",  (0,0), (-1,0), colors.white),
        ("FONTNAME",   (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE",   (0,0), (-1,-1), 7),
        ("GRID",       (0,0), (-1,-1), 0.4, colors.grey),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white, colors.HexColor("#eff6ff")]),
    ]))
    elems.append(tbl)
    doc.build(elems)
    buf.seek(0)
    return buf


def _build_xlsx(title: str, headers: list, rows: list) -> io.BytesIO:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = title[:31]
    ws.append(headers)
    for row in rows:
        ws.append([str(c) for c in row])
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def _build_docx(title: str, headers: list, rows: list) -> io.BytesIO:
    doc = DocxDoc()
    doc.add_heading(title, 0)
    doc.add_paragraph(datetime.now().strftime("Generated: %Y-%m-%d %H:%M"))
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.cell(0, i).text = h
    for row in rows:
        r = tbl.add_row()
        for i, val in enumerate(row):
            r.cells[i].text = str(val)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def run(payload: dict) -> dict:
    """
    payload keys:
      report_type: one of attendance | quiz | teaching_plan | question_bank
      formats: list of pdf | xlsx | docx
      title (optional): custom report title
    Returns:
      { report_type, links: { pdf: url, xlsx: url, ... } }
    """
    sheets      = SheetsService()
    report_type = payload.get("report_type", "attendance")
    formats     = payload.get("formats", ["pdf", "xlsx"])
    title       = payload.get("title", report_type.replace("_", " ").title())

    tab, headers = REPORT_CONFIG.get(report_type, ("attendance", []))
    records      = sheets.read_all(tab)
    rows         = [[str(r.get(h, "")) for h in headers] for r in records]

    ts    = datetime.now().strftime("%Y%m%d_%H%M")
    links = {}

    for fmt in formats:
        fname = f"{report_type}_{ts}.{fmt}"
        if fmt == "pdf":
            buf  = _build_pdf(title, headers, rows)
            mime = "application/pdf"
        elif fmt == "xlsx":
            buf  = _build_xlsx(title, headers, rows)
            mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        elif fmt == "docx":
            buf  = _build_docx(title, headers, rows)
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            continue
        links[fmt] = _upload_to_drive(fname, mime, buf)

    return {"report_type": report_type, "links": links}
